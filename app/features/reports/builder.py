import io

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.features.reports.schemas import ReportData


def build_report_document(report_data: ReportData) -> io.BytesIO:
    document = Document()

    _add_header(document, report_data)

    if report_data.summary is not None:
        _add_summary_section(document, report_data)

    if report_data.by_category is not None:
        _add_by_category_section(document, report_data)

    if report_data.by_wallet is not None:
        _add_by_wallet_section(document, report_data)

    if report_data.transactions is not None:
        _add_transactions_section(document, report_data)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _add_header(document: Document, report_data: ReportData) -> None:
    title = document.add_heading("Financial Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    period_paragraph = document.add_paragraph()
    period_paragraph.add_run(
        f"Period: {report_data.date_from.isoformat()} — {report_data.date_to.isoformat()}"
    )

    currency_paragraph = document.add_paragraph()
    currency_paragraph.add_run(f"Currency: {report_data.target_currency}")


def _add_summary_section(document: Document, report_data: ReportData) -> None:
    summary = report_data.summary
    document.add_heading("Summary", level=1)

    document.add_paragraph(f"Total income: {summary.total_income} {summary.currency}")
    document.add_paragraph(f"Total expense: {summary.total_expense} {summary.currency}")
    document.add_paragraph(f"Net: {summary.net} {summary.currency}")


def _add_by_category_section(document: Document, report_data: ReportData) -> None:
    document.add_heading("Expenses by Category", level=1)

    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Category"
    header_cells[1].text = "Total"
    header_cells[2].text = "Percentage"

    for item in report_data.by_category:
        row_cells = table.add_row().cells
        row_cells[0].text = item.category
        row_cells[1].text = str(item.total)
        row_cells[2].text = f"{item.percentage}%"


def _add_by_wallet_section(document: Document, report_data: ReportData) -> None:
    document.add_heading("Wallet Balances", level=1)

    note = document.add_paragraph()
    note_run = note.add_run("Balances shown are as of report generation time, not the selected period end date.")
    note_run.italic = True
    note_run.font.size = Pt(9)

    table = document.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Wallet"
    header_cells[1].text = "Balance"
    header_cells[2].text = "Currency"

    for item in report_data.by_wallet:
        row_cells = table.add_row().cells
        row_cells[0].text = item.wallet_name
        row_cells[1].text = str(item.balance)
        row_cells[2].text = item.currency


def _add_transactions_section(document: Document, report_data: ReportData) -> None:
    document.add_heading("Transactions", level=1)

    table = document.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"

    header_cells = table.rows[0].cells
    header_cells[0].text = "Date"
    header_cells[1].text = "Type"
    header_cells[2].text = "Description"
    header_cells[3].text = "Amount"
    header_cells[4].text = "Currency"

    for tx in report_data.transactions:
        row_cells = table.add_row().cells
        row_cells[0].text = tx.transaction_date.strftime("%Y-%m-%d %H:%M")
        row_cells[1].text = tx.operation_code
        row_cells[2].text = tx.description or "-"
        row_cells[3].text = str(tx.amount)
        row_cells[4].text = tx.currency